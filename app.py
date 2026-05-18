import os
import asyncio
import sys
from datetime import datetime
from fastapi import FastAPI, HTTPException, UploadFile, File, Form

if sys.stdout.encoding.lower() != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except AttributeError:
        # Fallback for older python or restricted environments
        pass
from pydantic import BaseModel
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import Optional, Dict

from crawlers.trend_crawler import run_trend_pipeline
from crawlers.science_crawler import run_science_pipeline
from crawlers.social_crawler import run_social_pipeline
from crawlers.article_scraper import scrape_sciencedaily_article
from crawlers.pdf_parser import extract_text_from_pdf, extract_text_from_bytes, UnreadablePDFError
from crawlers.science_crawler import extract_science_mechanisms
from engine.pipeline import run_v3_pipeline
from engine.script_generator import run_dual_agent_script_generation
import io
from docx import Document
from fastapi.responses import StreamingResponse
from db import store

app = FastAPI(title="雙軌科普腳本自動化引擎 V3", version="3.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/")
async def read_index():
    return FileResponse("static/index.html")

# ─── Data API (獲取資料庫內容) ───

@app.get("/api/data/trends")
def get_trends(page: int = 1, limit: int = 15):
    offset = (page - 1) * limit
    return store.fetch_latest_trends(limit=limit, offset=offset)

@app.get("/api/data/social")
def get_social(page: int = 1, limit: int = 15):
    offset = (page - 1) * limit
    return store.fetch_latest_social(limit=limit, offset=offset)

@app.get("/api/data/science")
def get_science(page: int = 1, limit: int = 15, search: Optional[str] = None):
    offset = (page - 1) * limit
    return store.fetch_latest_science(limit=limit, offset=offset, search=search)

@app.get("/api/history")
def get_history():
    return store.fetch_generation_history()

@app.delete("/api/data/{item_type}")
def delete_data_item(item_type: str, url: str):
    """刪除指定類型和 URL 的單筆資料。
    item_type: science | social | trend
    url: 要刪除的構圖
    """
    if item_type not in ("science", "social", "trend"):
        raise HTTPException(status_code=400, detail="無效的資料類型")
    success = store.delete_item(item_type, url)
    if not success:
        raise HTTPException(status_code=404, detail="找不到該筆資料")
    return {"status": "success", "message": "已成功刪除"}

# ─── Crawl API (觸發爬蟲) ───

@app.post("/api/crawl/trends")
def crawl_trends():
    run_trend_pipeline()
    return {"status": "success", "message": "時事爬取完成"}

@app.post("/api/crawl/social")
def crawl_social():
    run_social_pipeline()
    return {"status": "success", "message": "社群爬取完成"}

@app.post("/api/crawl/science")
def crawl_science():
    run_science_pipeline()
    return {"status": "success", "message": "科學爬取完成"}

@app.post("/api/crawl/all")
async def crawl_all():
    # 這裡可以改成 async/await 並發執行以加速
    run_trend_pipeline()
    run_social_pipeline()
    run_science_pipeline()
    return {"status": "success", "message": "全平台爬取完成"}


# ─── Generate API (觸發引擎配對與生成) ───

class GenerateRequest(BaseModel):
    locked_items: Optional[Dict[str, Optional[str]]] = None

@app.post("/api/generate")
async def generate_script(req: GenerateRequest):
    """
    執行 V3 流程：
    1. 呼叫 pipeline 進行 [科普核心] → [配對] → [3視角 Hook] → [Critic 審查]。
    """
    try:
        # 直接執行 V3 pipeline，傳入使用者鎖定的 URL 字典
        result = run_v3_pipeline(req.locked_items if req.locked_items else {})
        
        if result.get("error"):
            return {"status": "error", "message": result["error"]}

        quality_warning = ""
        if not result.get("passed", False):
            quality_warning = f"⚠️ 品質未達門檻（{result.get('critic_score', 0)}/{8}），以下為最佳結果"

        return {
            "status": "success",
            "science_core": result.get("science_core", "無科學核心內容"),
            "reasoning": result.get("reasoning", "無企劃屬性邏輯"),
            "mechanism": result.get("mechanism", "未知機制"),
            "hooks": result.get("hooks", ["", "", ""]),
            "hook_evaluations": result.get("hook_evaluations", []),
            "critic_score": result.get("critic_score", 0),
            "critic_breakdown": result.get("critic_breakdown", {}),
            "critic_comment": result.get("critic_comment", ""),
            "quality_warning": quality_warning,
            "matched_trend": result.get("matched_trend", {}),
            "matched_science": result.get("matched_science", {}),
            "matched_social": result.get("matched_social", {})
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"引擎內部錯誤：{str(e)}")


class BuildScriptRequest(BaseModel):
    science_url: str
    social_url: Optional[str] = ""
    hook_text: str


def _validate_pdf_content(science_text: str, web_article_text: str, science_url: str) -> str:
    """
    驗證下載的 PDF/全文是否為正確論文。
    比對策略（依優先順序）：
    1. 有 paper_doi → 用 regex 從 PDF 中擷取 DOI，精確比對
    2. 有 paper_title → 用關鍵字命中率比對（閾值 0.4）
    3. 兩者皆無 → 無法驗證，直接 fallback 到 web_article_text

    回傳：通過驗證的 science_text，或 fallback 的 web_article_text
    """
    import re

    # 從 DB 取得正確的 paper_doi 與 paper_title
    science_article = store.get_science_by_url(science_url)
    paper_doi = (science_article or {}).get("paper_doi", "").strip()
    paper_title = (science_article or {}).get("paper_title", "").strip()

    pdf_head = science_text[:5000]  # 只掃前 5000 字元（論文標題/摘要通常在開頭）
    print(f"[PDF 驗證] 開始驗證 PDF 正確性...")
    print(f"[PDF 驗證] 預期 DOI: {paper_doi or '(無)'} | 預期論文標題: {paper_title or '(無)'}")

    # === 情境 1：有 DOI → regex 精確比對 ===
    if paper_doi:
        doi_pattern = r'\b(10\.\d{4,9}/[-._;()/:a-zA-Z0-9]+)\b'
        found_dois = re.findall(doi_pattern, pdf_head, re.IGNORECASE)
        # 標準化比較（去掉結尾可能多出的標點）
        paper_doi_clean = paper_doi.strip(".")
        for found_doi in found_dois:
            if found_doi.strip(".").lower() == paper_doi_clean.lower():
                print(f"[PDF 驗證] DOI 比對通過: {found_doi}")
                return science_text
        if found_dois:
            print(f"[PDF 驗證] 警告：PDF 中找到 DOI {found_dois[0]}，但與預期 DOI {paper_doi} 不符！")
        else:
            # PDF 中找不到 DOI（例如 trafilatura 全文），降級到 paper_title 比對
            print(f"[PDF 驗證] PDF 中未找到 DOI，改用論文標題比對...")
            if paper_title:
                return _validate_by_title(science_text, web_article_text, paper_title, pdf_head)
            else:
                print(f"[PDF 驗證] 無 paper_title 可比對，fallback 到 ScienceDaily 介紹文章。")
                return web_article_text
        # DOI 存在但不匹配 → fallback
        print(f"[PDF 驗證] DOI 不匹配，fallback 到 ScienceDaily 介紹文章。")
        return web_article_text

    # === 情境 2：無 DOI，但有 paper_title ===
    if paper_title:
        return _validate_by_title(science_text, web_article_text, paper_title, pdf_head)

    # === 情境 3：兩者皆無 → 直接 fallback ===
    print(f"[PDF 驗證] 無 DOI 也無 paper_title，無法驗證 PDF 正確性，fallback 到 ScienceDaily 介紹文章。")
    return web_article_text


def _validate_by_title(science_text: str, web_article_text: str, paper_title: str, pdf_head: str, threshold: float = 0.4) -> str:
    """
    用 paper_title 關鍵字比對驗證 PDF 內容正確性。
    取標題中長度 > 4 的核心英文詞彙，計算在 PDF 開頭的命中率。
    """
    # 去掉常見停用詞，只保留有意義的名詞/動詞
    stopwords = {"with", "that", "this", "from", "into", "have", "been", "their",
                 "which", "through", "using", "after", "between", "during"}
    title_words = [w.lower() for w in paper_title.split() if len(w) > 4 and w.lower() not in stopwords]

    if not title_words:
        print(f"[PDF 驗證] paper_title 過短或全為停用詞，無法比對，fallback。")
        return web_article_text

    pdf_head_lower = pdf_head.lower()
    matched = sum(1 for w in title_words if w in pdf_head_lower)
    ratio = matched / len(title_words)
    print(f"[PDF 驗證] 標題關鍵字比對：{matched}/{len(title_words)} 命中，命中率 {ratio:.2f}（閾值 {threshold}）")

    if ratio >= threshold:
        print(f"[PDF 驗證] 標題比對通過，使用 PDF 內容。")
        return science_text
    else:
        print(f"[PDF 驗證] 標題比對失敗，PDF 可能下載錯誤，fallback 到 ScienceDaily 介紹文章。")
        return web_article_text

@app.post("/api/build_script")
async def build_script(req: BuildScriptRequest):
    """
    雙大腦生腳本流程：拉全文 -> 提取迷因 -> Agent 1 拆解加比喻 -> Agent 2 生成泛科學腳本
    """
    try:
        # 0. 即時萃取論文標題與 DOI (若資料庫尚無紀錄)
        science_article = store.get_science_by_url(req.science_url)
        if science_article and not science_article.get("paper_doi"):
            from crawlers.article_scraper import extract_paper_metadata
            meta = extract_paper_metadata(req.science_url)
            if meta.get("doi") or meta.get("paper_title"):
                store.update_science_paper_metadata(req.science_url, meta.get("paper_title", ""), meta.get("doi", ""))
                print(f"[build_script] 已即時擷取並儲存 DOI: {meta.get('doi')} / 標題: {meta.get('paper_title')}")
                
        # 1. 嘗試從 OpenAlex 自動下載 PDF (若本地無檔案)
        pdf_path = store.get_pdf_path_by_url(req.science_url)
        
        if not pdf_path or not os.path.exists(pdf_path):
            # 從 DB 讀取最新的 metadata (因為 Step 0 可能剛寫入 DOI)
            science_article = store.get_science_by_url(req.science_url)
            if science_article:
                paper_title = science_article.get("paper_title", "")
                paper_doi = science_article.get("paper_doi", "")
                if paper_title or paper_doi:
                    from crawlers.pdf_downloader import fetch_pdf_from_openalex
                    new_pdf_path = fetch_pdf_from_openalex(title=paper_title, doi=paper_doi)
                    if new_pdf_path:
                        store.update_science_pdf_path(req.science_url, new_pdf_path)
                        pdf_path = new_pdf_path
                        
        # 2. 抓取全文 — 優先使用 PDF (包含剛下載的)
        use_web_fallback = False
        if pdf_path and os.path.exists(pdf_path):
            print(f"[build_script] 使用本地 PDF: {pdf_path}")
            try:
                # 若 downloader 存的是 .txt（trafilatura 全文擷取），直接讀取
                if pdf_path.endswith(".txt"):
                    with open(pdf_path, "r", encoding="utf-8") as f:
                        science_text = f.read()
                    print(f"[build_script] 讀取 trafilatura 擷取的全文 ({len(science_text)} 字元)")
                    if not science_text:
                        use_web_fallback = True
                else:
                    science_text = extract_text_from_pdf(pdf_path)
                    if not science_text:
                        use_web_fallback = True
                        print("⚠️ 爬取 PDF 失敗（解析出空文字），轉向使用網頁原文擷取。")
            except UnreadablePDFError as e:
                use_web_fallback = True
                print(f"⚠️ 爬取 PDF 失敗（掃描檔或亂碼：{e}），轉向使用網頁原文擷取。")
        else:
            use_web_fallback = True
            print("⚠️ 爬取 PDF 失敗（可能受付費牆保護或無資源），轉向使用網頁原文擷取。")
            
        # 3. 無論如何都爬取網頁文章（作為新聞框架參考，或當作 Fallback）
        scrape_res = scrape_sciencedaily_article(req.science_url)
        if not scrape_res["success"]:
            raise HTTPException(status_code=400, detail=f"科學原稿爬取失敗: {scrape_res['error']}")
        
        web_article_text = scrape_res["text"]
        
        if use_web_fallback:
            science_text = web_article_text
        else:
            # === PDF 內容驗證：確認下載的 PDF 是否為正確論文 ===
            science_text = _validate_pdf_content(
                science_text=science_text,
                web_article_text=web_article_text,
                science_url=req.science_url
            )
        # 從資料庫撈出先前由 science_crawler 存入的正式標題，若無則降級使用當次爬蟲找到的標籤
        science_title = store.get_science_title_by_url(req.science_url)
        if not science_title:
            science_title = scrape_res.get("title", "未命名科學文章")
        
        # 2. 獲取迷因 Context
        meme_context = ""
        if req.social_url:
            meme_context = store.get_meme_context_by_url(req.social_url)
            
        # 3. 獲取腳本範本 (Style Reference)
        template_text = ""
        template_path = os.path.join(os.path.dirname(__file__), "腳本範本.md")
        if os.path.exists(template_path):
            with open(template_path, "r", encoding="utf-8") as f:
                template_text = f.read()
                
        # 4. 觸發雙 Agent 大腦
        script_markdown = run_dual_agent_script_generation(
            science_text=science_text,
            meme_context=meme_context,
            hook_text=req.hook_text,
            template_text=template_text,
            article_title=science_title,
            article_url=req.science_url,
            web_article_text=web_article_text
        )
        
        return {"status": "success", "script": script_markdown}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class ExportDocxRequest(BaseModel):
    markdown_text: str

@app.post("/api/export_docx")
async def export_docx(req: ExportDocxRequest):
    """將生成的 Markdown 轉換為 DOCX 供下載"""
    try:
        doc = Document()
        doc.add_heading('泛科學風格 - 自動生成腳本', 0)
        
        # 簡易的逐行解析 (簡單切割與加粗)
        for line in req.markdown_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('#'):
                level = line.count('#')
                text = line.replace('#', '').strip()
                doc.add_heading(text, level=min(level, 9))
            else:
                doc.add_paragraph(line)
                
        # 儲存到 BytesIO 並回傳
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return StreamingResponse(
            file_stream, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=script_generated.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ─── Save API (儲存至資料庫) ───

class SaveRequest(BaseModel):
    science_core: str
    mechanism: str
    hooks_json: Optional[str] = "[]"
    critic_score: Optional[int] = 0
    critic_breakdown_json: Optional[str] = "{}"
    matched_trend_url: Optional[str] = ""
    matched_science_url: Optional[str] = ""
    matched_social_url: Optional[str] = ""
    reasoning: Optional[str] = ""
    critic_comment: Optional[str] = ""
    hook_evaluations_json: Optional[str] = "[]"

@app.post("/api/save")
async def save_inspiration(req: SaveRequest):
    """將生成的靈感儲存至 SQLite history 表。"""
    try:
        if not req.science_core:
            raise HTTPException(status_code=400, detail="沒有可以保存的腳本！")

        store.save_history(req.dict())
        return {"status": "success", "message": "成功儲存靈感至資料庫"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"儲存失敗：{str(e)}")

# ─── Surprise API (隨機挑選科學文章生成) ───

@app.post("/api/surprise")
async def surprise_generate():
    """隨機從 DB 中挑選一篇科學文獻，自動配對時事/社群後生成 Hook。"""
    try:
        import random
        science_data = store.fetch_latest_science(limit=30).get("data", [])
        if not science_data:
            return {"status": "error", "message": "資料庫無科學文獻，請先執行爬取！"}
        
        chosen = random.choice(science_data)
        result = run_v3_pipeline({"science_url": chosen["url"]})
        
        if result.get("error"):
            return {"status": "error", "message": result["error"]}

        return {
            "status": "success",
            "science_core": result.get("science_core", ""),
            "reasoning": result.get("reasoning", "無企劃屬性邏輯"),
            "mechanism": result.get("mechanism", ""),
            "hooks": result.get("hooks", ["", "", ""]),
            "hook_evaluations": result.get("hook_evaluations", []),
            "critic_score": result.get("critic_score", 0),
            "critic_breakdown": result.get("critic_breakdown", {}),
            "critic_comment": result.get("critic_comment", ""),
            "matched_science": result.get("matched_science", {}),
            "matched_trend": result.get("matched_trend", {}),
            "matched_social": result.get("matched_social", {})
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Surprise 生成失敗：{str(e)}")


# ─── Science Evaluate API (評估科學文章可信度) ───

class EvaluateRequest(BaseModel):
    url: str

@app.post("/api/science/evaluate")
async def evaluate_science(req: EvaluateRequest):
    """根據 URL 回傳該科學來源的可信度星級與白/黑名單狀態。"""
    from config import SCIENTIFIC_WHITELIST, SCIENTIFIC_BLACKLIST, CREDIBILITY_SCORES
    
    url_lower = req.url.lower()
    
    # 黑名單檢查
    is_blacklisted = any(bl in url_lower for bl in SCIENTIFIC_BLACKLIST)
    
    # 白名單檢查
    is_whitelisted = any(wl in url_lower for wl in SCIENTIFIC_WHITELIST)
    
    # 可信度分數
    score = 1
    for domain, s in CREDIBILITY_SCORES.items():
        if domain != "default" and domain in url_lower:
            score = s
            break

    return {
        "url": req.url,
        "credibility_score": score,
        "is_whitelisted": is_whitelisted,
        "is_blacklisted": is_blacklisted,
        "stars": "★" * score + "☆" * (3 - score)
    }


# ─── Manual Upload API (手動新增資料) ───

@app.post("/api/manual/science")
async def manual_add_science(
    title: str = Form(...),
    url: str = Form(""),
    abstract: str = Form(""),
    pdf_file: UploadFile = File(None)
):
    """手動新增科學論文至資料庫。PDF 存本地，abstract 用 LLM 提取 mechanism/category。"""
    import time as _time
    try:
        # 處理 URL — 若留空自動生成
        if not url.strip():
            url = f"manual://science-{int(_time.time())}"

        # 處理 PDF 上傳
        pdf_path = ""
        if pdf_file and pdf_file.filename:
            content = await pdf_file.read()
            
            try:
                # 預先檢查 PDF 是否可讀
                extract_text_from_bytes(content)
            except UnreadablePDFError as e:
                return {"status": "unreadable_pdf", "message": str(e)}
                
            os.makedirs(os.path.join(os.path.dirname(__file__), "data", "pdfs"), exist_ok=True)
            safe_name = f"{int(_time.time())}_{pdf_file.filename}"
            save_path = os.path.join(os.path.dirname(__file__), "data", "pdfs", safe_name)
            with open(save_path, "wb") as f:
                f.write(content)
            pdf_path = save_path
            print(f"[Manual] PDF saved to {save_path}")

        # 用 abstract 提取 mechanism + category
        article_data = {
            "title": title,
            "summary": abstract or "(No abstract provided)",
            "url": url,
            "source": "手動上傳",
            "pipeline": "Manual",
            "credibility_score": 1,
            "pdf_path": pdf_path
        }

        # 若有 abstract，嘗試用 LLM 提取 mechanism / category
        if abstract.strip():
            try:
                enriched = extract_science_mechanisms([article_data])
                if enriched:
                    article_data = enriched[0]
            except Exception as e:
                print(f"⚠️ 手動上傳 LLM 提取失敗 (非致命): {e}")
                article_data["mechanism"] = ""
                article_data["category"] = ""
        else:
            article_data["mechanism"] = ""
            article_data["category"] = ""

        result = store.save_science([article_data])
        return {"status": "success", "message": f"科學論文已新增", **result}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"新增失敗：{str(e)}")


class ManualSocialRequest(BaseModel):
    title: str
    summary: str
    url: Optional[str] = ""
    category: Optional[str] = "meme"

@app.post("/api/manual/social")
async def manual_add_social(req: ManualSocialRequest):
    """手動新增迷因/社群內容至資料庫。"""
    import time as _time
    try:
        item_url = req.url.strip() if req.url else f"manual://social-{int(_time.time())}"

        item = {
            "title": req.title,
            "summary": req.summary,
            "url": item_url,
            "source": "手動上傳",
            "category": req.category or "meme",
            "is_short": False,
            "mechanism": "",
            "matched_keywords": [],
            "thumbnail": ""
        }

        result = store.save_social([item])
        return {"status": "success", "message": "迷因已新增", **result}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"新增失敗：{str(e)}")


if __name__ == "__main__":
    import uvicorn
    print("啟動 V3 引擎伺服器...")
    uvicorn.run("app:app", host="127.0.0.1", port=8000, reload=True)
